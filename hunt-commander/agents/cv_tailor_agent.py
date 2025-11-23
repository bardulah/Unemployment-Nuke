"""CV Tailor agent for customizing CVs for specific job postings."""
import os
from typing import Dict, Any
from pathlib import Path
from datetime import datetime
from openai import OpenAI
from anthropic import Anthropic
from docx import Document
from .base_agent import BaseAgent

class CVTailorAgent(BaseAgent):
    """Agent responsible for tailoring CVs to match specific job requirements."""

    def __init__(self, config: Dict[str, Any]):
        """Initialize the CV tailor agent.

        Args:
            config: Configuration dictionary
        """
        super().__init__(config)
        self.openai_client = None
        self.anthropic_client = None
        self.model = config.get('cv_tailor_model', 'gpt-4')
        self.max_pages = config.get('cv_max_pages', 2)
        self.output_dir = Path(config.get('output_directory', 'data/cv/tailored'))
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._initialize_clients()

    def _initialize_clients(self):
        """Initialize AI clients if API keys are available."""
        openai_key = os.getenv('OPENAI_API_KEY')
        anthropic_key = os.getenv('ANTHROPIC_API_KEY')

        if openai_key:
            self.openai_client = OpenAI(api_key=openai_key)
            self.log_info("OpenAI client initialized")

        if anthropic_key:
            self.anthropic_client = Anthropic(api_key=anthropic_key)
            self.log_info("Anthropic client initialized")

    def execute(
        self,
        job: Dict[str, Any],
        cv_content: str,
        cv_template_path: str = None
    ) -> Dict[str, Any]:
        """Tailor CV for a specific job posting.

        Args:
            job: Job dictionary with requirements
            cv_content: Original CV content
            cv_template_path: Path to CV template file (optional)

        Returns:
            Dictionary with tailored CV information
        """
        self.log_info(f"Tailoring CV for: {job.get('title', 'Unknown')} at {job.get('company', 'Unknown')}")

        # Generate tailored CV content
        tailored_content = self._generate_tailored_cv(job, cv_content)

        if not tailored_content:
            self.log_error("Failed to generate tailored CV")
            return None

        # Save tailored CV
        cv_file = self._save_tailored_cv(job, tailored_content, cv_template_path)

        result = {
            'job_title': job.get('title'),
            'company': job.get('company'),
            'cv_file_path': str(cv_file),
            'tailored_content': tailored_content,
            'created_at': datetime.now().isoformat()
        }

        self.log_info(f"Tailored CV saved to: {cv_file}")
        return result

    def _generate_tailored_cv(self, job: Dict[str, Any], cv_content: str) -> str:
        """Generate tailored CV content using AI.

        Args:
            job: Job dictionary
            cv_content: Original CV content

        Returns:
            Tailored CV content as string
        """
        # Use AI for intelligent CV tailoring
        if 'claude' in self.model.lower() and self.anthropic_client:
            return self._ai_tailor_anthropic(job, cv_content)
        elif self.openai_client:
            return self._ai_tailor_openai(job, cv_content)
        else:
            self.log_error("No AI client available for CV tailoring")
            return cv_content  # Return original if no AI available

    def _ai_tailor_anthropic(self, job: Dict[str, Any], cv_content: str) -> str:
        """Use Anthropic Claude for CV tailoring.

        Args:
            job: Job dictionary
            cv_content: Original CV content

        Returns:
            Tailored CV content
        """
        prompt = self._build_tailoring_prompt(job, cv_content)

        try:
            response = self.anthropic_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )

            tailored_content = response.content[0].text
            return tailored_content

        except Exception as e:
            self.log_error(f"Error in AI CV tailoring: {e}")
            return cv_content

    def _ai_tailor_openai(self, job: Dict[str, Any], cv_content: str) -> str:
        """Use OpenAI GPT for CV tailoring.

        Args:
            job: Job dictionary
            cv_content: Original CV content

        Returns:
            Tailored CV content
        """
        prompt = self._build_tailoring_prompt(job, cv_content)

        try:
            response = self.openai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert CV writer who tailors resumes to match specific job requirements."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7
            )

            tailored_content = response.choices[0].message.content
            return tailored_content

        except Exception as e:
            self.log_error(f"Error in AI CV tailoring: {e}")
            return cv_content

    def _build_tailoring_prompt(self, job: Dict[str, Any], cv_content: str) -> str:
        """Build the prompt for CV tailoring.

        Args:
            job: Job dictionary
            cv_content: Original CV content

        Returns:
            Prompt string
        """
        prompt = f"""You are an expert CV/resume writer. Your task is to tailor the following CV to perfectly match the job requirements below.

JOB POSTING:
Title: {job.get('title', 'N/A')}
Company: {job.get('company', 'N/A')}
Location: {job.get('location', 'N/A')}
Description: {job.get('description', 'N/A')}
Requirements: {job.get('requirements', 'N/A')}
Full Description: {job.get('full_description', 'N/A')}

ORIGINAL CV:
{cv_content}

Please create a tailored version of this CV that:

1. HIGHLIGHTS relevant experience and skills that match the job requirements
2. REORDERS sections to put the most relevant information first
3. ADJUSTS language to mirror keywords from the job description (without lying)
4. EMPHASIZES achievements relevant to this specific role
5. ADDS or expands on relevant skills/technologies mentioned in the job posting
6. KEEPS it concise (maximum {self.max_pages} pages worth of content)
7. MAINTAINS all factual information - DO NOT fabricate experience or skills
8. USES strong action verbs and quantifiable achievements where possible

Important rules:
- DO NOT invent skills, experience, or qualifications
- DO NOT change dates or job titles
- DO change emphasis, ordering, and descriptions to better match the role
- DO use keywords from the job description naturally
- DO make it ATS (Applicant Tracking System) friendly

Format the output as a clean, well-structured CV in markdown format with clear sections:
- Contact Information
- Professional Summary (tailored to this role)
- Technical Skills (prioritized by relevance)
- Work Experience (detailed and tailored)
- Education
- Additional sections as relevant (Certifications, Projects, etc.)

Begin the tailored CV now:"""

        return prompt

    def _save_tailored_cv(
        self,
        job: Dict[str, Any],
        tailored_content: str,
        cv_template_path: str = None
    ) -> Path:
        """Save tailored CV to file.

        Args:
            job: Job dictionary
            tailored_content: Tailored CV content
            cv_template_path: Path to template file (optional)

        Returns:
            Path to saved CV file
        """
        # Create filename
        company = job.get('company', 'Unknown').replace(' ', '_').replace('/', '_')
        title = job.get('title', 'Position').replace(' ', '_').replace('/', '_')
        timestamp = datetime.now().strftime('%Y%m%d')
        filename = f"CV_{company}_{title}_{timestamp}.md"

        cv_file = self.output_dir / filename

        # Save as markdown
        with open(cv_file, 'w', encoding='utf-8') as f:
            f.write(tailored_content)

        # Also try to save as DOCX if template is provided
        if cv_template_path and os.path.exists(cv_template_path):
            try:
                docx_file = cv_file.with_suffix('.docx')
                self._create_docx_from_markdown(tailored_content, docx_file)
                self.log_info(f"Also created DOCX version: {docx_file}")
            except Exception as e:
                self.log_error(f"Failed to create DOCX version: {e}")

        return cv_file

    def _create_docx_from_markdown(self, markdown_content: str, output_path: Path):
        """Create a DOCX file from markdown content.

        Args:
            markdown_content: Markdown formatted CV
            output_path: Path to save DOCX file

        Note:
            This is a simple conversion. For better formatting,
            consider using pandoc or python-markdown
        """
        doc = Document()

        # Simple markdown to docx conversion
        lines = markdown_content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                continue

            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Numbered lists
            elif line and line[0].isdigit() and '. ' in line:
                doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
            # Regular paragraphs
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
